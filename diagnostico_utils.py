import os
import json
import tempfile
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from ga4_utils import (
    fetch_ga4_kpis,
    fetch_funil_conversao,
    fetch_produtos_mais_vendidos,
    fetch_categorias_mais_vendidas,
    fetch_tecnologia_usuarios,
    fetch_regioes_mais_acessadas,
    fetch_engajamento_site,
    fetch_paginas_mais_acessadas,
    fetch_conversoes_por_canal,
    fetch_funil_abandono
)

load_dotenv()

# --- Cliente Groq com endpoint compatível OpenAI ---
client = OpenAI(
    api_key=st.secrets["GROQ_API_KEY"],
    base_url="https://api.groq.com/openai/v1"
)

# ------------------------------
# Utilidades de modelo (Groq)
# ------------------------------
def _listar_modelos_disponiveis():
    """Tenta listar modelos do provedor; retorna nomes (str)."""
    try:
        modelos = client.models.list()
        return [m.id for m in getattr(modelos, "data", []) if hasattr(m, "id")]
    except Exception:
        # Sem permissão/listagem indisponível – segue com fallback estático
        return []

def _sequencia_modelos():
    """
    Monta a sequência de tentativas de modelos:
    1) GROQ_MODEL do secrets (se existir)
    2) Modelos retornados por /models priorizando Llama recentes
    3) Lista de candidatos estática (ampla)
    """
    preferido = st.secrets.get("GROQ_MODEL", "").strip()
    ordem = []

    if preferido:
        ordem.append(preferido)

    # 2) prioriza modelos ativos retornados pelo provedor
    ativos = _listar_modelos_disponiveis()
    # Heurística de prioridade: Llama mais novos/maiores primeiro
    prioridade = (
        "llama-3.3-70b", "llama-3.3-8b",
        "llama-3.2-90b", "llama-3.2-70b", "llama-3.2-11b", "llama-3.2-8b",
        "llama-3.1-70b", "llama-3.1-8b",
        "llama3-70b", "llama3-8b"
    )
    ativos_ordenados = sorted(
        [m for m in ativos if any(m.startswith(p) for p in prioridade)],
        key=lambda m: next((i for i, p in enumerate(prioridade) if m.startswith(p)), 999)
    )
    ordem.extend([m for m in ativos_ordenados if m not in ordem])

    # 3) fallback amplo (inclui variações comuns no Groq)
    candidatos = [
        # Llama 3.3 (novos)
        "llama-3.3-70b-versatile",
        "llama-3.3-70b-specdec",
        "llama-3.3-8b-instant",
        # Llama 3.2
        "llama-3.2-90b-text-preview",
        "llama-3.2-11b-text-preview",
        "llama-3.2-8b-text-preview",
        # Llama 3.1 (alguns workspaces ainda possuem)
        "llama-3.1-70b-versatile",
        "llama-3.1-8b-instant",
        # Llama 3.0 (legado – alguns ambientes ainda expõem)
        "llama3-70b-8192",
        "llama3-8b-8192",
    ]
    for c in candidatos:
        if c not in ordem:
            ordem.append(c)

    # remove duplicados preservando ordem
    visto = set()
    final = []
    for m in ordem:
        if m and m not in visto:
            final.append(m)
            visto.add(m)
    return final

# ------------------------------
# Coleta de dados
# ------------------------------
@st.cache_data(ttl=3600)
def coletar_dados_dashboard(property_id, start_date, end_date, customer_root):
    kpis = fetch_ga4_kpis(property_id, start_date, end_date, customer_root)
    funil = fetch_funil_conversao(property_id, start_date, end_date, customer_root)
    produtos = fetch_produtos_mais_vendidos(property_id, start_date, end_date, customer_root)
    categorias = fetch_categorias_mais_vendidas(property_id, start_date, end_date, customer_root)
    df_disp, _ = fetch_tecnologia_usuarios(property_id, start_date, end_date, customer_root)
    regioes = fetch_regioes_mais_acessadas(property_id, start_date, end_date, customer_root)
    engajamento = fetch_engajamento_site(property_id, start_date, end_date, customer_root)
    paginas = fetch_paginas_mais_acessadas(property_id, start_date, end_date, customer_root)
    canais = fetch_conversoes_por_canal(property_id, start_date, end_date, customer_root)
    abandono = fetch_funil_abandono(property_id, start_date, end_date, customer_root)

    return {
        "kpis": kpis,
        "funil": funil,
        "produtos": produtos.to_dict(orient="records") if not produtos.empty else [],
        "categorias": categorias.to_dict(orient="records") if not categorias.empty else [],
        "dispositivo_destaque": df_disp.sort_values("Sessões", ascending=False).iloc[0]["Categoria"] if not df_disp.empty else "n/d",
        "cidade_top": regioes.sort_values("Acessos", ascending=False).iloc[0]["Cidade"] if not regioes.empty else "n/d",
        "engajamento": engajamento.iloc[-1]["Taxa de Engajamento (%)"] if not engajamento.empty else 0,
        "sessões_totais": int(engajamento["Acessos Totais"].sum()) if not engajamento.empty else 0,
        "paginas_top": paginas["Página"].head(3).tolist() if not paginas.empty else [],
        "paginas": paginas.to_dict(orient="records") if not paginas.empty else [],
        "canais": canais.to_dict(orient="records") if not canais.empty else [],
        "abandono": abandono,
        "periodo": f"{start_date} a {end_date}",
        "ponto_forte": max(produtos.to_dict(orient="records"), key=lambda x: x.get("Receita (R$)", 0), default=None) if not produtos.empty else None,
        "tendencia_negativa": "Conversão abaixo de 0.5%" if kpis.get("taxa_conversao", 1) < 0.005 else ""
    }

# ------------------------------
# Prompt
# ------------------------------
def gerar_prompt(dados, cliente, inicio, fim):
    kpis = dados["kpis"]
    canal = dados.get("top_canal", "N/D")
    produto = dados.get("top_produto", "N/D")
    regiao = dados.get("top_regiao", "N/D")

    return f"""
Você é um analista sênior de Web Analytics especializado em e-commerce.

Abaixo estão os dados do cliente **{cliente}**, no período de **{inicio} a {fim}**:

---
📊 **KPIs Gerais**
- Receita Total: R$ {kpis['receita_total']:,.2f}
- Vendas: {kpis['vendas']}
- Taxa de Conversão: {kpis['taxa_conversao']:.2%}
- Ticket Médio: R$ {kpis['ticket_medio']:,.2f}

🏆 **Destaques**
- Canal com maior receita: {canal}
- Produto mais vendido: {produto}
- Região com mais acessos: {regiao}

---
🎯 **Instruções para você (IA):**
1. Analise os dados e destaque **pontos de atenção** e **boas oportunidades**.
2. Faça **perguntas estratégicas** que ajudariam o cliente a refletir sobre o desempenho.
3. Dê sugestões práticas de otimização para tráfego, conversão ou engajamento.
4. Evite repetir os dados e foque na **interpretação inteligente**.
5. Seja consultivo, direto e claro.

Escreva como um consultor de performance digital falando com um gestor de e-commerce.
"""

# ------------------------------
# Chamada à IA (com auto-descoberta e fallbacks)
# ------------------------------
def _extrair_mensagem_erro(e: Exception) -> str:
    """
    Tenta extrair uma mensagem útil do erro do SDK compatível.
    """
    try:
        # openai-python geralmente expõe e.response ou e.body
        body = getattr(e, "body", None) or {}
        if isinstance(body, dict) and "error" in body:
            msg = body["error"].get("message") or body["error"]
            return str(msg)
        # fallback para a string do próprio erro
        return str(e)
    except Exception:
        return str(e)

def chamar_ia(prompt: str) -> str:
    modelos = _sequencia_modelos()
    ultima_excecao = None

    for modelo in modelos:
        try:
            resp = client.chat.completions.create(
                model=modelo,
                messages=[
                    {"role": "system", "content": "Você é um analista de dados Web Analytics consultivo e especialista em e-commerce."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=900  # ajustável conforme necessidade
            )
            return resp.choices[0].message.content
        except Exception as e:
            ultima_excecao = e
            import traceback
            st.warning(f"Falha ao chamar o modelo '{modelo}'. Tentando o próximo…")
            # Mensagem do provedor (útil p/ saber se foi 'decommissioned', quota, etc.)
            st.caption("Detalhes (dev):")
            st.code(_extrair_mensagem_erro(e))

    # Se todos os modelos falharem, levanta o último erro para o Streamlit capturar
    raise ultima_excecao

# ------------------------------
# Exportadores
# ------------------------------
def exportar_txt(texto):
    return texto.encode("utf-8")

def exportar_docx(texto):
    doc = Document()
    doc.add_heading("📋 Diagnóstico Automatizado de Performance GA4", level=1)
    for linha in texto.split("\n"):
        if linha.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(linha)
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    return temp.name
